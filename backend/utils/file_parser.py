"""
文件内容提取工具
支持：PDF / Word(.docx) / 纯文本(.txt/.md)
"""
import io


def extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return _parse_pdf(content)
    elif ext == "docx":
        return _parse_docx(content)
    elif ext in ("txt", "md", "markdown"):
        return content.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"不支持的文件格式：.{ext}，请上传 PDF / Word / txt / md")


def _parse_pdf(content: bytes) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(stream=content, filetype="pdf")
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts).strip()


def _parse_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # 表格内容也提取
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()
